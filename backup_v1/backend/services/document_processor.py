"""
Document Processor Service
===========================
Handles all document-related operations including parsing, processing,
preview generation, and final document creation.

This module provides:
- DOCX file parsing with structure preservation
- Content extraction from paragraphs and tables
- HTML preview generation with placeholder highlighting
- Final document generation with filled values

Author: Legal Tech Solutions
Date: October 2025
Version: 1.0.0
"""

import os
import re
import copy
import logging
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE

# Configure logging
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    A comprehensive document processor for legal documents.
    
    This class handles:
    - Parsing DOCX files while preserving formatting
    - Extracting text content and structure
    - Generating HTML previews with placeholder highlighting
    - Creating final documents with filled values
    """
    
    def __init__(self):
        """
        Initialize the DocumentProcessor with default settings.
        """
        self.placeholder_pattern = re.compile(r'\{\{([^}]+)\}\}')
        self.supported_formats = ['.docx', '.doc']
        logger.info("DocumentProcessor initialized")
    
    def parse_document(self, filepath: str) -> Dict[str, Any]:
        """
        Parse a DOCX document and extract its content structure.
        
        This method extracts:
        - All paragraphs with formatting information
        - Tables with cell data
        - Document metadata
        - Raw text for analysis
        
        Args:
            filepath (str): Path to the DOCX file to parse
            
        Returns:
            Dict[str, Any]: Document content structure with paragraphs, tables, and metadata
            
        Raises:
            FileNotFoundError: If the document file doesn't exist
            Exception: For document parsing errors
        """
        try:
            # Validate file exists
            if not os.path.exists(filepath):
                logger.error(f"Document file not found: {filepath}")
                raise FileNotFoundError(f"Document not found: {filepath}")
            
            # Open and parse the document
            doc = Document(filepath)
            logger.info(f"Opened document: {filepath}")
            
            # Initialize content structure
            content = {
                'paragraphs': [],
                'tables': [],
                'raw_text': '',
                'metadata': {
                    'sections': len(doc.sections),
                    'paragraphs_count': len(doc.paragraphs),
                    'tables_count': len(doc.tables),
                    'core_properties': self._extract_core_properties(doc)
                }
            }
            
            # Extract paragraphs with formatting
            full_text = []
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                
                # Skip empty paragraphs
                if not para_text:
                    continue
                
                # Extract paragraph data
                para_data = {
                    'index': i,
                    'text': para_text,
                    'style': paragraph.style.name if paragraph.style else 'Normal',
                    'alignment': str(paragraph.alignment) if paragraph.alignment else 'LEFT',
                    'runs': []
                }
                
                # Extract run-level formatting for precise reconstruction
                for run in paragraph.runs:
                    run_data = {
                        'text': run.text,
                        'bold': run.bold if run.bold is not None else False,
                        'italic': run.italic if run.italic is not None else False,
                        'underline': run.underline if run.underline else False,
                        'font_size': run.font.size.pt if run.font.size else 12,
                        'font_name': run.font.name if run.font.name else 'Calibri',
                        'font_color': self._get_color_value(run.font.color) if run.font.color else None
                    }
                    para_data['runs'].append(run_data)
                
                content['paragraphs'].append(para_data)
                full_text.append(para_text)
            
            # Extract tables with structure preservation
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    'index': table_idx,
                    'rows': [],
                    'dimensions': (len(table.rows), len(table.columns) if table.rows else 0),
                    'style': table.style.name if table.style else None
                }
                
                # Extract each cell's content
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        cell_data = {
                            'text': cell_text,
                            'row': row_idx,
                            'col': cell_idx,
                            'paragraphs': [p.text for p in cell.paragraphs]
                        }
                        row_data.append(cell_data)
                        
                        # Add cell text to full text for searching
                        if cell_text:
                            full_text.append(cell_text)
                    
                    table_data['rows'].append(row_data)
                
                content['tables'].append(table_data)
            
            # Combine all text for analysis
            content['raw_text'] = '\n'.join(full_text)
            
            logger.info(f"Successfully parsed document with {len(content['paragraphs'])} paragraphs and {len(content['tables'])} tables")
            return content
            
        except FileNotFoundError:
            raise
        except Exception as e:
            logger.error(f"Error parsing document {filepath}: {str(e)}")
            raise Exception(f"Document parsing failed: {str(e)}")
    
    def _extract_core_properties(self, doc: Document) -> Dict[str, Any]:
        """
        Extract document metadata and properties.
        
        Args:
            doc (Document): python-docx Document object
            
        Returns:
            Dict containing document properties
        """
        try:
            core_props = doc.core_properties
            return {
                'title': core_props.title or 'Untitled',
                'author': core_props.author or 'Unknown',
                'created': str(core_props.created) if core_props.created else None,
                'modified': str(core_props.modified) if core_props.modified else None,
                'subject': core_props.subject or None,
                'keywords': core_props.keywords or None
            }
        except:
            return {}
    
    def _get_color_value(self, color) -> Optional[str]:
        """
        Extract color value from docx color object.
        
        Args:
            color: docx color object
            
        Returns:
            Hex color string or None
        """
        try:
            if color.rgb:
                return f"#{color.rgb}"
            return None
        except:
            return None
    
    def generate_preview(self, content: Dict[str, Any], placeholders: List[Dict],
                        filled_values: Dict[str, str], current_index: Optional[int] = None) -> str:
        """
        Generate an HTML preview of the document with highlighted placeholders.
        
        This creates an HTML representation showing:
        - Current field highlighted in RED (with field name)
        - Filled placeholders in GREEN (with field name)
        - Unfilled placeholders without highlight (only when not current)
        - Document structure and formatting
        
        Args:
            content (Dict): Document content structure
            placeholders (List[Dict]): List of detected placeholders
            filled_values (Dict[str, str]): Dictionary of filled placeholder values
            current_index (Optional[int]): Index of current field being filled (None if none)
            
        Returns:
            str: HTML string representing the document preview
        """
        try:
            preview_html = ['<div class="document-preview">']
            
            # Add custom CSS for preview styling
            preview_html.append("""
            <style>
                .document-preview {
                    font-family: 'Calibri', 'Arial', sans-serif;
                    line-height: 1.6;
                    color: #333;
                    background: white;
                    padding: 40px;
                    max-width: 800px;
                    margin: 0 auto;
                }
                .paragraph {
                    margin-bottom: 12px;
                }
                .heading {
                    font-weight: bold;
                    font-size: 1.2em;
                    margin-top: 20px;
                    margin-bottom: 10px;
                }
                .title {
                    font-weight: bold;
                    font-size: 1.5em;
                    margin-bottom: 20px;
                    text-align: center;
                }
                .placeholder-current {
                    background-color: #fee2e2;
                    color: #991b1b;
                    padding: 4px 8px;
                    border-radius: 4px;
                    font-weight: 700;
                    border: 2px solid #dc2626;
                    box-shadow: 0 0 8px rgba(220, 38, 38, 0.3);
                    display: inline-block;
                }
                .placeholder-filled {
                    background-color: #d4edda;
                    color: #155724;
                    padding: 2px 6px;
                    border-radius: 3px;
                    font-weight: 600;
                    border: 1px solid #c3e6cb;
                    cursor: pointer;
                }
                .placeholder-filled:hover {
                    background-color: #c3e6cb;
                    text-decoration: underline;
                }
                .placeholder-unfilled {
                    background-color: transparent;
                    color: #333;
                    padding: 2px 6px;
                    border-radius: 3px;
                    font-weight: 500;
                    border: 1px dashed #ccc;
                    opacity: 0.6;
                }
                .document-table {
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                }
                .document-table td, .document-table th {
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }
                .document-table th {
                    background-color: #f3f4f6;
                    font-weight: bold;
                }
                .document-table tr:nth-child(even) {
                    background-color: #ffffff;
                }
            </style>
            """)
            
            # Process paragraphs
            for para in content.get('paragraphs', []):
                para_text = para['text']
                para_location = para['index']  # Get paragraph location
                
                # CRITICAL FIX: Filter placeholders that belong to THIS specific paragraph
                # This prevents replacing placeholders from other locations
                para_placeholders = [
                    (idx, p) for idx, p in enumerate(placeholders)
                    if p.get('location_type') == 'paragraph' and p.get('location') == para_location
                ]
                
                # Track what we've replaced to avoid double-replacement within same paragraph
                replaced_in_this_para = set()
                
                # Replace placeholders with highlighted versions (only for THIS paragraph)
                for idx, placeholder in para_placeholders:
                    placeholder_key = placeholder['key']
                    placeholder_text = placeholder['original']
                    placeholder_name = placeholder.get('name', 'Field')
                    placeholder_id = placeholder.get('id', placeholder_key)  # Get ID early
                    
                    # Check if this placeholder pattern exists in this paragraph
                    if placeholder_text not in para_text:
                        continue
                    
                    # Check if we already replaced this pattern in this paragraph
                    # (handles edge case where same pattern appears multiple times in one paragraph)
                    if placeholder_text in replaced_in_this_para:
                        continue
                    
                    is_current = (current_index is not None and idx == current_index)
                    # Check filled status by ID first, then key as fallback
                    is_filled = placeholder_id in filled_values or placeholder_key in filled_values
                    
                    # Determine CSS class and styling
                    if is_current:
                        # Current field - highlight in RED with field name
                        field_name_display = f"[{placeholder_name}]"
                        replacement = f'''<span class="placeholder-current" 
                                       title="Field: {placeholder_name} - Currently filling this field"
                                       data-ph="{placeholder_id}"
                                       data-key="{placeholder_key}"
                                       data-index="{idx}">
                                       {field_name_display}</span>'''
                    elif is_filled:
                        # Filled field - show in GREEN with value and field name
                        # Look up by ID first, then key as fallback
                        field_value = filled_values.get(placeholder_id, filled_values.get(placeholder_key, ''))
                        replacement = f'''<span class="placeholder-filled" 
                                       title="Field: {placeholder_name} - Click to edit (Value: {field_value})"
                                       data-ph="{placeholder_id}"
                                       data-key="{placeholder_key}"
                                       data-index="{idx}">
                                       {field_value}</span>'''
                    else:
                        # Unfilled field - show minimally (no highlight until it's current)
                        replacement = f'''<span class="placeholder-unfilled" 
                                       title="Field: {placeholder_name} - Not yet filled"
                                       data-ph="{placeholder_id}"
                                       data-key="{placeholder_key}"
                                       data-index="{idx}">
                                       {placeholder_text}</span>'''
                    
                    # Replace only the FIRST occurrence
                    para_text = para_text.replace(placeholder_text, replacement, 1)
                    replaced_in_this_para.add(placeholder_text)
                
                # Apply paragraph styling based on style name
                style_class = 'paragraph'
                if 'heading' in para['style'].lower():
                    style_class = 'heading'
                elif 'title' in para['style'].lower():
                    style_class = 'title'
                
                # Escape any remaining HTML entities
                para_text = self._escape_html(para_text, preserve_spans=True)
                
                preview_html.append(f'<p class="{style_class}">{para_text}</p>')
            
            # Process tables
            for table in content.get('tables', []):
                preview_html.append('<table class="document-table">')
                
                for row_idx, row in enumerate(table['rows']):
                    preview_html.append('<tr>')
                    
                    for cell in row:
                        cell_text = cell['text'] if isinstance(cell, dict) else cell
                        
                        # Apply placeholder highlighting in table cells
                        for placeholder in placeholders:
                            placeholder_id = placeholder.get('id', placeholder['key'])
                            # Check by ID first, then key as fallback
                            if placeholder_id in filled_values or placeholder['key'] in filled_values:
                                value = filled_values.get(placeholder_id, filled_values.get(placeholder['key'], ''))
                                cell_text = cell_text.replace(
                                    placeholder['original'],
                                    f'<span class="placeholder-filled">{value}</span>'
                                )
                            else:
                                cell_text = cell_text.replace(
                                    placeholder['original'],
                                    f'<span class="placeholder-unfilled">{placeholder["original"]}</span>'
                                )
                        
                        # Use th for first row (header)
                        tag = 'th' if row_idx == 0 else 'td'
                        preview_html.append(f'<{tag}>{cell_text}</{tag}>')
                    
                    preview_html.append('</tr>')
                
                preview_html.append('</table>')
            
            preview_html.append('</div>')
            
            return '\n'.join(preview_html)
            
        except Exception as e:
            logger.error(f"Error generating preview: {str(e)}")
            # Return a simple error message in preview
            return f'<div class="document-preview"><p style="color: red;">Error generating preview: {str(e)}</p></div>'
    
    def _escape_html(self, text: str, preserve_spans: bool = False) -> str:
        """
        Escape HTML entities while optionally preserving span tags.
        
        Args:
            text (str): Text to escape
            preserve_spans (bool): Whether to preserve span tags
            
        Returns:
            str: Escaped text
        """
        if not preserve_spans:
            text = text.replace('&', '&amp;')
            text = text.replace('<', '&lt;')
            text = text.replace('>', '&gt;')
            text = text.replace('"', '&quot;')
            text = text.replace("'", '&#39;')
        
        return text
    
    def generate_final_document(self, template_path: str, output_path: str,
                               placeholders: List[Dict], filled_values: Dict[str, str]) -> bool:
        """
        Generate the final document with all placeholders replaced.
        
        This method:
        1. Opens the template document
        2. Replaces all placeholders with filled values
        3. Highlights the replaced text (optional)
        4. Saves the completed document
        
        Args:
            template_path (str): Path to the template document
            output_path (str): Path where the final document should be saved
            placeholders (List[Dict]): List of placeholders to replace
            filled_values (Dict[str, str]): Dictionary of values to insert
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Validate inputs
            if not os.path.exists(template_path):
                logger.error(f"Template document not found: {template_path}")
                return False
            
            # Ensure output directory exists
            output_dir = os.path.dirname(output_path)
            Path(output_dir).mkdir(parents=True, exist_ok=True)
            
            # Load the template document
            doc = Document(template_path)
            logger.info(f"Loaded template document: {template_path}")
            
            # Track replacements made
            replacements_made = 0
            
            # NOTE: Removed auto-fill logic - context-aware deduplication in placeholder_detector.py
            # now handles duplicate $[___] patterns correctly by differentiating based on context
            # (Purchase Amount vs Valuation Cap). Each field gets its own value from filled_values.
            
            # Process all paragraphs - match by ORIGINAL document indices
            for i, paragraph in enumerate(doc.paragraphs):
                if not paragraph.text.strip():
                    continue  # Skip empty but keep index counting
                
                paragraph_text = paragraph.text
                text_changed = False
                
                # Filter placeholders that belong to THIS paragraph
                para_placeholders = [
                    p for p in placeholders 
                    if p.get('location_type') == 'paragraph' and p.get('location') == i
                ]
                
                # Replace each placeholder with its specific value
                for placeholder in para_placeholders:
                    placeholder_id = placeholder.get('id', placeholder['key'])
                    # Check by ID first, then key as fallback
                    if (placeholder_id in filled_values or placeholder['key'] in filled_values) and placeholder['original'] in paragraph_text:
                        old_value = placeholder['original']
                        new_value = filled_values.get(placeholder_id, filled_values.get(placeholder['key'], ''))
                        
                        # Replace only this occurrence
                        paragraph_text = paragraph_text.replace(old_value, new_value, 1)
                        text_changed = True
                        replacements_made += 1
                        
                        logger.info(f"Replaced '{placeholder['name']}' (id: {placeholder_id}) in para {i}: '{old_value}' → '{new_value}'")
                
                # Handle auto-fill for header Company Name
                if '[Company Name]' in paragraph_text:
                    for ph in placeholders:
                        ph_id = ph.get('id', ph['key'])
                        if ph['name'] == 'Company Name' and (ph_id in filled_values or ph['key'] in filled_values):
                            value = filled_values.get(ph_id, filled_values.get(ph['key'], ''))
                            paragraph_text = paragraph_text.replace('[Company Name]', value, 1)
                            text_changed = True
                            replacements_made += 1
                            break
                
                # Update paragraph if text changed
                if text_changed:
                    # Clear existing runs and add new text as single run
                    paragraph.clear()
                    paragraph.add_run(paragraph_text)
            
            # Process all tables
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        # Process each paragraph in the cell
                        for paragraph in cell.paragraphs:
                            # Replace at paragraph level to handle placeholders split across runs
                            paragraph_text = paragraph.text
                            text_changed = False
                            
                            # Build location key for this table cell
                            table_location = f"{table_idx}-{row_idx}-{col_idx}"
                            
                            # Replace placeholders that match THIS table cell location
                            for placeholder in placeholders:
                                # Check if this placeholder belongs to this table cell
                                placeholder_id = placeholder.get('id', placeholder['key'])
                                if (placeholder.get('location_type') == 'table' and
                                    str(placeholder.get('location')) == table_location and
                                    placeholder['original'] in paragraph_text and 
                                    (placeholder_id in filled_values or placeholder['key'] in filled_values)):
                                    
                                    value = filled_values.get(placeholder_id, filled_values.get(placeholder['key'], ''))
                                    paragraph_text = paragraph_text.replace(
                                        placeholder['original'],
                                        value,
                                        1  # Replace only first occurrence
                                    )
                                    replacements_made += 1
                                    text_changed = True
                                    logger.debug(f"Replaced {placeholder['name']} (id: {placeholder_id}) in table {table_location}")
                            
                            # Only update if text changed
                            if text_changed:
                                # Clear existing runs and add new text as single run
                                paragraph.clear()
                                paragraph.add_run(paragraph_text)
            
            # Save the completed document
            doc.save(output_path)
            logger.info(f"Successfully generated document with {replacements_made} replacements: {output_path}")
            
            return True
            
        except Exception as e:
            logger.error(f"Error generating final document: {str(e)}")
            return False
    
    def validate_document_structure(self, filepath: str) -> Dict[str, Any]:
        """
        Validate document structure and identify potential issues.
        
        This method checks for:
        - Empty documents
        - Extremely large documents
        - Complex tables
        - Corrupted files
        
        Args:
            filepath (str): Path to the document to validate
            
        Returns:
            Dict[str, Any]: Validation results with issues and warnings
        """
        try:
            validation_result = {
                'valid': True,
                'issues': [],
                'warnings': [],
                'stats': {}
            }
            
            # Check file exists
            if not os.path.exists(filepath):
                validation_result['valid'] = False
                validation_result['issues'].append('File does not exist')
                return validation_result
            
            # Check file size
            file_size_mb = os.path.getsize(filepath) / (1024 * 1024)
            validation_result['stats']['file_size_mb'] = round(file_size_mb, 2)
            
            if file_size_mb > 10:
                validation_result['valid'] = False
                validation_result['issues'].append(f'File too large: {file_size_mb:.1f}MB (max 10MB)')
                return validation_result
            
            # Try to open the document
            try:
                doc = Document(filepath)
            except:
                validation_result['valid'] = False
                validation_result['issues'].append('Unable to open document - file may be corrupted')
                return validation_result
            
            # Count elements
            num_paragraphs = len(doc.paragraphs)
            num_tables = len(doc.tables)
            num_sections = len(doc.sections)
            
            validation_result['stats'] = {
                'paragraphs': num_paragraphs,
                'tables': num_tables,
                'sections': num_sections,
                'file_size_mb': round(file_size_mb, 2)
            }
            
            # Check for empty document
            if num_paragraphs == 0 and num_tables == 0:
                validation_result['valid'] = False
                validation_result['issues'].append('Document is empty')
            
            # Check for extremely large documents
            if num_paragraphs > 1000:
                validation_result['warnings'].append(f'Document is very large ({num_paragraphs} paragraphs) and may take time to process')
            
            # Check for complex tables
            for table in doc.tables:
                if len(table.rows) > 50:
                    validation_result['warnings'].append('Document contains very large tables (50+ rows)')
                    break
                if len(table.columns) > 10:
                    validation_result['warnings'].append('Document contains very wide tables (10+ columns)')
                    break
            
            # Check for nested tables (not supported well)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if len(cell.tables) > 0:
                            validation_result['warnings'].append('Document contains nested tables which may not display correctly')
                            break
            
            return validation_result
            
        except Exception as e:
            logger.error(f"Error validating document: {str(e)}")
            return {
                'valid': False,
                'issues': [f'Validation error: {str(e)}'],
                'warnings': [],
                'stats': {}
            }
    
    def extract_text_only(self, filepath: str) -> str:
        """
        Extract plain text from a document without formatting.
        
        Args:
            filepath (str): Path to the document
            
        Returns:
            str: Plain text content of the document
        """
        try:
            doc = Document(filepath)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            return ""